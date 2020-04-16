import os
from docx import Document
from messages import Messages
from files import Files

class Documentation:
    new_documentation=True
    doc_type=""
    code_root_dir=""
    target_root_dir=""
    filename=""
    document=""
    def __init__(self,*args, **kwargs):
        doc_type=kwargs['doc_type']
        code_root_dir=kwargs['code_root_dir']
        target_root_dir=kwargs['target_root_dir']
        filename=kwargs['filename']
        if(Files.dir_exists(target_root_dir) and Files.file_exists(str(target_root_dir)+"/"+str(filename),False)==False):
            Messages.success("Starting documentation")
            document=Document()
            document.add_heading("Heading one")
            document.add_paragraph("Paragraph one")
            document.save(target_root_dir+"/"+str(filename))
        else:
            Messages.error("Check whether directory "+str(target_root_dir)+" exists")


    def check_existing_doc(self):
        pass 

    def new_word_document(self):
        pass

    def resume_documentation(self):
        pass

    def documentation_version(self):
        pass

class Comments:

    def filter_comments(self,filepath,delimeter):
        pass

    def filter_text(self,filepath):
        pass

    def comment_text(self,comments):
        pass 

    def code_blue_print(self,fileComments):
        pass

    def read_code_file(self,path):
        pass

    def filter_comments(self,filedata):
        pass 

    def code_segment(self):
        pass 

    def group_paragraph(self):
        pass 

D=Documentation(doc_type="Full",code_root_dir="/home/antony/Pit/Projects/desktop/autoDoc",target_root_dir="/home/antony/Desktop",filename="docone.docx")